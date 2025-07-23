#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Enhanced File utilities module for RAG Document Indexer
Handles advanced document parsing with specialized support for Word documents, 
image extraction, and hybrid text+image processing (English documents only)
"""

import os
import io
import zipfile
import tempfile
from pathlib import Path
from datetime import datetime
from llama_index.core import SimpleDirectoryReader, Document

# Advanced document parsing imports
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("INFO: python-docx not available. Install with: pip install python-docx")

try:
    import pypandoc
    PANDOC_AVAILABLE = True
except ImportError:
    PANDOC_AVAILABLE = False
    print("INFO: pypandoc not available. Install with: pip install pypandoc")

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False


def clean_content_from_null_bytes(content):
    """
    Clean content from null bytes and other problematic characters
    
    Args:
        content: Text content to clean
    
    Returns:
        str: Cleaned content
    """
    if not isinstance(content, str):
        return content
    
    # Remove null bytes (\u0000) and other problematic characters
    content = content.replace('\u0000', '').replace('\x00', '').replace('\x01', '').replace('\x02', '')
    
    # Remove control characters (except newlines and tabs)
    cleaned_content = ''.join(char for char in content 
                            if ord(char) >= 32 or char in '\n\t\r')
    
    return cleaned_content


def clean_metadata_recursive(obj):
    """
    Recursively clean metadata from null bytes
    
    Args:
        obj: Object to clean (dict, list, str, etc.)
    
    Returns:
        Cleaned object
    """
    if isinstance(obj, dict):
        return {k: clean_metadata_recursive(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [clean_metadata_recursive(v) for v in obj]
    elif isinstance(obj, str):
        # Remove null bytes and limit string length
        cleaned = obj.replace('\u0000', '').replace('\x00', '')
        return cleaned[:1000]  # Limit metadata string length
    else:
        return obj


def safe_read_file(file_path, max_size=50*1024*1024):
    """
    Safely read file with basic error handling for English files
    
    Args:
        file_path: Path to the file to read
        max_size: Maximum file size in bytes (default 50MB)
    
    Returns:
        tuple: (content, error_code) where content is file text or None if failed
    """
    try:
        # Check file size first
        file_size = os.path.getsize(file_path)
        if file_size > max_size:
            return None, "FILE_TOO_LARGE"
        
        if file_size == 0:
            return None, "EMPTY_FILE"
        
        # Try UTF-8 first (standard for English)
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                # Clean null bytes
                content = clean_content_from_null_bytes(content)
                if content.strip():
                    return content, None
        except UnicodeDecodeError:
            # Fallback to Latin-1 for older English files
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
                    # Clean null bytes
                    content = clean_content_from_null_bytes(content)
                    if content.strip():
                        return content, "LATIN1_FALLBACK"
            except Exception:
                pass
        
        # Last resort: binary mode with forced conversion
        try:
            with open(file_path, 'rb') as f:
                raw_content = f.read()
                content = raw_content.decode('utf-8', errors='replace')
                content = content.replace('\ufffd', ' ')  # Remove replacement chars
                content = ''.join(c for c in content if c.isprintable() or c.isspace())
                
                # Clean null bytes
                content = clean_content_from_null_bytes(content)
                
                if content.strip():
                    return content, "FORCED_DECODE"
                
        except Exception:
            pass
        
        return None, "NO_READABLE_CONTENT"
        
    except Exception as e:
        return None, f"FATAL_ERROR_{type(e).__name__}"


class AdvancedDocxParser:
    """Advanced parser for .docx files with image extraction and structure preservation"""
    
    def __init__(self, extract_images=True, preserve_structure=True, extract_tables=True):
        """
        Initialize advanced DOCX parser
        
        Args:
            extract_images: Whether to extract images from documents
            preserve_structure: Whether to preserve document structure (headers, etc.)
            extract_tables: Whether to extract table content
        """
        self.extract_images = extract_images
        self.preserve_structure = preserve_structure
        self.extract_tables = extract_tables
        self.is_available = DOCX_AVAILABLE and PIL_AVAILABLE
        
        if not self.is_available:
            print("WARNING: Advanced DOCX parsing not available. Missing dependencies.")
    
    def extract_images_from_docx(self, docx_file_path):
        """
        Extract all images from DOCX file
        
        Args:
            docx_file_path: Path to DOCX file
        
        Returns:
            list: List of (image_data, image_name, image_format) tuples
        """
        if not self.is_available or not self.extract_images:
            return []
        
        images = []
        
        try:
            # Open DOCX as ZIP file to access images
            with zipfile.ZipFile(docx_file_path, 'r') as docx_zip:
                # Find all image files in the ZIP
                image_files = [f for f in docx_zip.namelist() 
                              if f.startswith('word/media/') and 
                              any(f.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp'])]
                
                for i, image_file in enumerate(image_files):
                    try:
                        # Extract image data
                        image_data = docx_zip.read(image_file)
                        
                        # Determine image format
                        image_name = os.path.basename(image_file)
                        image_format = os.path.splitext(image_name)[1].lower()
                        
                        # Create a more descriptive name
                        base_name = os.path.splitext(os.path.basename(docx_file_path))[0]
                        enhanced_name = f"{base_name}_image_{i+1}{image_format}"
                        
                        images.append((image_data, enhanced_name, image_format))
                        
                    except Exception as e:
                        print(f"   WARNING: Failed to extract image {image_file}: {e}")
                
            if images:
                print(f"   INFO: Extracted {len(images)} images from {os.path.basename(docx_file_path)}")
                
        except Exception as e:
            print(f"   ERROR: Failed to extract images from {docx_file_path}: {e}")
        
        return images
    
    def extract_table_content(self, doc):
        """
        Extract content from all tables in the document
        
        Args:
            doc: python-docx Document object
        
        Returns:
            str: Formatted table content
        """
        if not self.extract_tables:
            return ""
        
        table_content = []
        
        try:
            for i, table in enumerate(doc.tables):
                table_text = f"\n--- Table {i+1} ---\n"
                
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        # Clean cell text and remove extra whitespace
                        cell_text = ' '.join(cell.text.strip().split())
                        row_cells.append(cell_text)
                    
                    # Join cells with tab separator
                    table_text += '\t'.join(row_cells) + '\n'
                
                table_text += "--- End Table ---\n"
                table_content.append(table_text)
                
        except Exception as e:
            print(f"   WARNING: Failed to extract table content: {e}")
        
        return '\n'.join(table_content)
    
    def extract_structured_content(self, doc):
        """
        Extract content while preserving document structure
        
        Args:
            doc: python-docx Document object
        
        Returns:
            tuple: (structured_text, metadata_info)
        """
        if not self.preserve_structure:
            # Simple text extraction
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            return '\n'.join(full_text), {}
        
        structured_content = []
        metadata_info = {
            'headings': [],
            'paragraph_count': 0,
            'heading_count': 0,
            'list_count': 0
        }
        
        try:
            for paragraph in doc.paragraphs:
                if not paragraph.text.strip():
                    continue
                
                # Check if paragraph is a heading
                if paragraph.style.name.startswith('Heading'):
                    level = paragraph.style.name.replace('Heading ', '')
                    heading_marker = '#' * min(int(level) if level.isdigit() else 1, 6)
                    structured_content.append(f"{heading_marker} {paragraph.text.strip()}")
                    
                    metadata_info['headings'].append({
                        'level': level,
                        'text': paragraph.text.strip()
                    })
                    metadata_info['heading_count'] += 1
                
                # Check if paragraph is a list item
                elif paragraph.style.name.startswith('List'):
                    structured_content.append(f"• {paragraph.text.strip()}")
                    metadata_info['list_count'] += 1
                
                # Regular paragraph
                else:
                    structured_content.append(paragraph.text.strip())
                    metadata_info['paragraph_count'] += 1
            
        except Exception as e:
            print(f"   WARNING: Failed to extract structured content: {e}")
            # Fallback to simple extraction
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            return '\n'.join(full_text), {'extraction_fallback': True}
        
        return '\n'.join(structured_content), metadata_info
    
    def parse_docx_file(self, file_path):
        """
        Parse DOCX file with advanced features
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            tuple: (main_document, extracted_images, parsing_info)
        """
        if not self.is_available:
            print(f"   WARNING: Advanced DOCX parsing not available for {os.path.basename(file_path)}")
            return None, [], {'error': 'dependencies_missing'}
        
        try:
            print(f"   INFO: Advanced parsing of {os.path.basename(file_path)}")
            
            # Open document
            doc = DocxDocument(file_path)
            
            # Extract structured text content
            structured_text, structure_metadata = self.extract_structured_content(doc)
            
            # Extract table content
            table_content = self.extract_table_content(doc)
            
            # Combine text and table content
            full_text = structured_text
            if table_content:
                full_text += f"\n\n{table_content}"
            
            # Clean text
            full_text = clean_content_from_null_bytes(full_text)
            
            # Extract images
            extracted_images = self.extract_images_from_docx(file_path)
            
            # Create enhanced metadata
            parsing_info = {
                'parser_type': 'advanced_docx',
                'structure_preserved': self.preserve_structure,
                'images_extracted': len(extracted_images),
                'tables_found': len(doc.tables) if hasattr(doc, 'tables') else 0,
                'structure_metadata': structure_metadata,
                'extraction_features': {
                    'extract_images': self.extract_images,
                    'preserve_structure': self.preserve_structure,
                    'extract_tables': self.extract_tables
                }
            }
            
            # Create main document
            clean_file_path = clean_content_from_null_bytes(str(file_path))
            clean_file_name = clean_content_from_null_bytes(os.path.basename(file_path))
            
            main_metadata = {
                'file_path': clean_file_path,
                'file_name': clean_file_name,
                'file_type': 'docx',
                'file_size': os.path.getsize(file_path),
                'parsing_info': parsing_info,
                'content_length': len(full_text),
                'extraction_method': 'advanced_docx_parser'
            }
            
            # Clean metadata
            clean_main_metadata = clean_metadata_recursive(main_metadata)
            
            main_document = Document(
                text=full_text,
                metadata=clean_main_metadata
            )
            
            print(f"   SUCCESS: Extracted {len(full_text)} characters, {len(extracted_images)} images")
            
            return main_document, extracted_images, parsing_info
            
        except Exception as e:
            print(f"   ERROR: Failed to parse DOCX file {file_path}: {e}")
            return None, [], {'error': str(e), 'parser_type': 'advanced_docx'}


class LegacyDocConverter:
    """Converter for legacy .doc files using pandoc or other methods"""
    
    def __init__(self):
        """Initialize legacy DOC converter"""
        self.pandoc_available = PANDOC_AVAILABLE
        
        if not self.pandoc_available:
            print("INFO: pandoc not available for .doc conversion. .doc files will use fallback methods.")
    
    def convert_doc_to_text(self, doc_file_path):
        """
        Convert .doc file to text using available methods
        
        Args:
            doc_file_path: Path to .doc file
        
        Returns:
            tuple: (text_content, conversion_info)
        """
        conversion_info = {'method': 'none', 'success': False}
        
        try:
            # Method 1: Try pandoc if available
            if self.pandoc_available:
                try:
                    text_content = pypandoc.convert_file(doc_file_path, 'plain')
                    text_content = clean_content_from_null_bytes(text_content)
                    
                    if text_content and text_content.strip():
                        conversion_info = {
                            'method': 'pandoc',
                            'success': True,
                            'content_length': len(text_content)
                        }
                        print(f"   INFO: Converted .doc using pandoc: {len(text_content)} characters")
                        return text_content, conversion_info
                
                except Exception as e:
                    print(f"   WARNING: Pandoc conversion failed: {e}")
            
            # Method 2: Try to read as binary and extract readable text (very basic)
            try:
                with open(doc_file_path, 'rb') as f:
                    raw_content = f.read()
                
                # Very basic text extraction from binary data
                # This is a fallback method and may not work well
                text_content = raw_content.decode('latin-1', errors='ignore')
                
                # Filter out non-printable characters and keep only readable text
                readable_chars = []
                for char in text_content:
                    if char.isprintable() or char.isspace():
                        readable_chars.append(char)
                
                filtered_content = ''.join(readable_chars)
                
                # Remove excessive whitespace and clean up
                lines = [line.strip() for line in filtered_content.split('\n') if line.strip()]
                final_content = '\n'.join(lines)
                
                # Clean null bytes
                final_content = clean_content_from_null_bytes(final_content)
                
                if final_content and len(final_content) > 100:  # Minimum reasonable length
                    conversion_info = {
                        'method': 'binary_extraction',
                        'success': True,
                        'content_length': len(final_content),
                        'warning': 'Low quality extraction method'
                    }
                    print(f"   WARNING: Using basic binary extraction for .doc file: {len(final_content)} characters")
                    return final_content, conversion_info
                
            except Exception as e:
                print(f"   WARNING: Binary extraction failed: {e}")
            
            # If all methods fail
            conversion_info = {
                'method': 'failed',
                'success': False,
                'error': 'No conversion method succeeded'
            }
            
            return "", conversion_info
            
        except Exception as e:
            conversion_info = {
                'method': 'error',
                'success': False,
                'error': str(e)
            }
            return "", conversion_info


class HybridDocumentProcessor:
    """Processor that combines text extraction with image OCR for complete document processing"""
    
    def __init__(self, config=None):
        """
        Initialize hybrid document processor
        
        Args:
            config: Configuration object with processing settings
        """
        self.config = config
        
        # Load settings from config
        if config:
            doc_settings = config.get_document_parsing_settings()
            self.advanced_parsing_enabled = doc_settings.get('advanced_parsing_enabled', True)
            self.extract_images = doc_settings.get('extract_images', True)
            self.preserve_structure = doc_settings.get('preserve_structure', True)
            self.extract_tables = doc_settings.get('extract_tables', True)
            self.hybrid_processing = doc_settings.get('hybrid_processing', True)
            self.combine_results = doc_settings.get('combine_results', True)
            self.image_quality = doc_settings.get('image_quality', 'high')
        else:
            # Default settings
            self.advanced_parsing_enabled = True
            self.extract_images = True
            self.preserve_structure = True
            self.extract_tables = True
            self.hybrid_processing = True
            self.combine_results = True
            self.image_quality = 'high'
        
        # Initialize specialized parsers
        self.docx_parser = AdvancedDocxParser(
            extract_images=self.extract_images,
            preserve_structure=self.preserve_structure,
            extract_tables=self.extract_tables
        ) if self.advanced_parsing_enabled else None
        
        self.doc_converter = LegacyDocConverter() if self.advanced_parsing_enabled else None
        
        # OCR processor will be injected when needed
        self.ocr_processor = None
    
    def set_ocr_processor(self, ocr_processor):
        """
        Set OCR processor for image processing
        
        Args:
            ocr_processor: OCR processor instance
        """
        self.ocr_processor = ocr_processor
    
    def process_docx_file(self, file_path):
        """
        Process DOCX file with advanced parsing and image extraction
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            list: List of Document objects (main document + image documents)
        """
        if not self.docx_parser or not self.advanced_parsing_enabled:
            # Fallback to simple processing
            return self._simple_file_processing(file_path)
        
        try:
            # Parse DOCX with advanced features
            main_document, extracted_images, parsing_info = self.docx_parser.parse_docx_file(file_path)
            
            documents = []
            
            # Add main document if successful
            if main_document:
                documents.append(main_document)
            
            # Process extracted images with OCR if enabled
            if extracted_images and self.hybrid_processing and self.ocr_processor:
                print(f"   INFO: Processing {len(extracted_images)} extracted images with OCR")
                
                for image_data, image_name, image_format in extracted_images:
                    try:
                        # Create temporary file for OCR processing
                        with tempfile.NamedTemporaryFile(suffix=image_format, delete=False) as temp_file:
                            temp_file.write(image_data)
                            temp_file_path = temp_file.name
                        
                        # Process with OCR
                        ocr_document = self.ocr_processor.process_single_image(temp_file_path)
                        
                        if ocr_document:
                            # Enhance metadata to indicate it's from a document
                            ocr_document.metadata.update({
                                'source_document': os.path.basename(file_path),
                                'extraction_method': 'docx_image_ocr',
                                'original_image_name': image_name
                            })
                            documents.append(ocr_document)
                        
                        # Clean up temporary file
                        os.unlink(temp_file_path)
                        
                    except Exception as e:
                        print(f"   WARNING: Failed to process extracted image {image_name}: {e}")
            
            return documents
            
        except Exception as e:
            print(f"   ERROR: Advanced DOCX processing failed for {file_path}: {e}")
            return self._simple_file_processing(file_path)
    
    def process_doc_file(self, file_path):
        """
        Process legacy .doc file
        
        Args:
            file_path: Path to .doc file
        
        Returns:
            list: List of Document objects
        """
        if not self.doc_converter or not self.advanced_parsing_enabled:
            return self._simple_file_processing(file_path)
        
        try:
            # Convert .doc to text
            text_content, conversion_info = self.doc_converter.convert_doc_to_text(file_path)
            
            if conversion_info['success'] and text_content:
                # Create document
                clean_file_path = clean_content_from_null_bytes(str(file_path))
                clean_file_name = clean_content_from_null_bytes(os.path.basename(file_path))
                
                metadata = {
                    'file_path': clean_file_path,
                    'file_name': clean_file_name,
                    'file_type': 'doc',
                    'file_size': os.path.getsize(file_path),
                    'conversion_info': conversion_info,
                    'content_length': len(text_content),
                    'extraction_method': 'legacy_doc_converter'
                }
                
                # Clean metadata
                clean_metadata = clean_metadata_recursive(metadata)
                
                document = Document(
                    text=text_content,
                    metadata=clean_metadata
                )
                
                return [document]
            else:
                print(f"   WARNING: Failed to convert .doc file: {file_path}")
                return self._simple_file_processing(file_path)
                
        except Exception as e:
            print(f"   ERROR: Legacy .doc processing failed for {file_path}: {e}")
            return self._simple_file_processing(file_path)
    
    def _simple_file_processing(self, file_path):
        """
        Fallback to simple file processing
        
        Args:
            file_path: Path to file
        
        Returns:
            list: List of Document objects
        """
        try:
            content, error_code = safe_read_file(file_path)
            
            if content:
                clean_file_path = clean_content_from_null_bytes(str(file_path))
                clean_file_name = clean_content_from_null_bytes(os.path.basename(file_path))
                
                metadata = {
                    'file_path': clean_file_path,
                    'file_name': clean_file_name,
                    'file_type': os.path.splitext(file_path)[1].lower(),
                    'file_size': os.path.getsize(file_path),
                    'extraction_method': 'simple_text_reader',
                    'content_length': len(content)
                }
                
                if error_code:
                    metadata['encoding_fallback'] = error_code
                
                # Clean metadata
                clean_metadata = clean_metadata_recursive(metadata)
                
                document = Document(
                    text=content,
                    metadata=clean_metadata
                )
                
                return [document]
            else:
                return []
                
        except Exception as e:
            print(f"   ERROR: Simple file processing failed for {file_path}: {e}")
            return []


class AdvancedDirectoryLoader:
    """
    Advanced directory loader with specialized document parsing and hybrid processing
    """
    
    def __init__(self, input_dir, recursive=True, config=None):
        """
        Initialize advanced directory loader
        
        Args:
            input_dir: Input directory path
            recursive: Whether to scan recursively
            config: Configuration object
        """
        self.input_dir = input_dir
        self.recursive = recursive
        self.config = config
        
        # Initialize hybrid processor
        self.hybrid_processor = HybridDocumentProcessor(config)
        
        # Statistics
        self.loading_stats = {
            'total_files_found': 0,
            'docx_files': 0,
            'doc_files': 0,
            'other_files': 0,
            'documents_created': 0,
            'images_extracted': 0,
            'processing_errors': 0,
            'advanced_parsing_used': 0,
            'fallback_used': 0
        }
    
    def set_ocr_processor(self, ocr_processor):
        """
        Set OCR processor for image processing
        
        Args:
            ocr_processor: OCR processor instance
        """
        self.hybrid_processor.set_ocr_processor(ocr_processor)
    
    def scan_files(self):
        """
        Scan directory for files and categorize them
        
        Returns:
            dict: Categorized file lists
        """
        file_categories = {
            'docx_files': [],
            'doc_files': [],
            'other_files': []
        }
        
        try:
            if self.recursive:
                file_iterator = Path(self.input_dir).rglob('*')
            else:
                file_iterator = Path(self.input_dir).glob('*')
            
            for file_path in file_iterator:
                if file_path.is_file():
                    file_ext = file_path.suffix.lower()
                    
                    if file_ext == '.docx':
                        file_categories['docx_files'].append(str(file_path))
                    elif file_ext == '.doc':
                        file_categories['doc_files'].append(str(file_path))
                    else:
                        file_categories['other_files'].append(str(file_path))
                    
                    self.loading_stats['total_files_found'] += 1
            
            self.loading_stats['docx_files'] = len(file_categories['docx_files'])
            self.loading_stats['doc_files'] = len(file_categories['doc_files'])
            self.loading_stats['other_files'] = len(file_categories['other_files'])
            
        except Exception as e:
            print(f"ERROR: Failed to scan directory {self.input_dir}: {e}")
        
        return file_categories
    
    def load_data(self):
        """
        Load and process all documents using advanced parsing
        
        Returns:
            tuple: (documents, loading_stats, failed_files)
        """
        print("?? Advanced Document Loading with Hybrid Processing")
        
        # Scan files
        file_categories = self.scan_files()
        
        if self.loading_stats['total_files_found'] == 0:
            print("WARNING: No files found in directory")
            return [], self.loading_stats, []
        
        print(f"   Found: {self.loading_stats['docx_files']} .docx, "
              f"{self.loading_stats['doc_files']} .doc, "
              f"{self.loading_stats['other_files']} other files")
        
        all_documents = []
        failed_files = []
        
        # Process DOCX files with advanced parsing
        if file_categories['docx_files']:
            print(f"   Processing {len(file_categories['docx_files'])} DOCX files...")
            for docx_file in file_categories['docx_files']:
                try:
                    documents = self.hybrid_processor.process_docx_file(docx_file)
                    if documents:
                        all_documents.extend(documents)
                        self.loading_stats['documents_created'] += len(documents)
                        self.loading_stats['advanced_parsing_used'] += 1
                        
                        # Count extracted images
                        for doc in documents:
                            if doc.metadata.get('extraction_method') == 'docx_image_ocr':
                                self.loading_stats['images_extracted'] += 1
                    else:
                        failed_files.append(docx_file)
                        self.loading_stats['processing_errors'] += 1
                        
                except Exception as e:
                    print(f"   ERROR: Failed to process {docx_file}: {e}")
                    failed_files.append(docx_file)
                    self.loading_stats['processing_errors'] += 1
        
        # Process DOC files with legacy conversion
        if file_categories['doc_files']:
            print(f"   Processing {len(file_categories['doc_files'])} DOC files...")
            for doc_file in file_categories['doc_files']:
                try:
                    documents = self.hybrid_processor.process_doc_file(doc_file)
                    if documents:
                        all_documents.extend(documents)
                        self.loading_stats['documents_created'] += len(documents)
                        self.loading_stats['advanced_parsing_used'] += 1
                    else:
                        failed_files.append(doc_file)
                        self.loading_stats['processing_errors'] += 1
                        
                except Exception as e:
                    print(f"   ERROR: Failed to process {doc_file}: {e}")
                    failed_files.append(doc_file)
                    self.loading_stats['processing_errors'] += 1
        
        # Process other files with standard SimpleDirectoryReader
        if file_categories['other_files']:
            print(f"   Processing {len(file_categories['other_files'])} other files...")
            
            # Create temporary directory with only other files for SimpleDirectoryReader
            other_documents = []
            for other_file in file_categories['other_files']:
                try:
                    # Use simple file reading for non-Word documents
                    documents = self.hybrid_processor._simple_file_processing(other_file)
                    if documents:
                        other_documents.extend(documents)
                        self.loading_stats['documents_created'] += len(documents)
                        self.loading_stats['fallback_used'] += 1
                    else:
                        failed_files.append(other_file)
                        self.loading_stats['processing_errors'] += 1
                        
                except Exception as e:
                    print(f"   ERROR: Failed to process {other_file}: {e}")
                    failed_files.append(other_file)
                    self.loading_stats['processing_errors'] += 1
            
            all_documents.extend(other_documents)
        
        # Final statistics
        total_success = len(all_documents)
        total_failed = len(failed_files)
        success_rate = (total_success / (total_success + total_failed) * 100) if (total_success + total_failed) > 0 else 0
        
        self.loading_stats.update({
            'total_documents_created': total_success,
            'total_failed': total_failed,
            'success_rate': success_rate
        })
        
        print(f"   ? Successfully processed: {total_success} documents")
        print(f"   ? Failed: {total_failed} files")
        print(f"   ?? Success rate: {success_rate:.1f}%")
        if self.loading_stats['images_extracted'] > 0:
            print(f"   ???  Images extracted and processed: {self.loading_stats['images_extracted']}")
        
        return all_documents, self.loading_stats, failed_files
    
    def get_processing_summary(self):
        """
        Get comprehensive processing summary
        
        Returns:
            dict: Processing summary with detailed statistics
        """
        return {
            'total_files_found': self.loading_stats['total_files_found'],
            'file_breakdown': {
                'docx_files': self.loading_stats['docx_files'],
                'doc_files': self.loading_stats['doc_files'], 
                'other_files': self.loading_stats['other_files']
            },
            'processing_results': {
                'documents_created': self.loading_stats['documents_created'],
                'images_extracted': self.loading_stats['images_extracted'],
                'processing_errors': self.loading_stats['processing_errors'],
                'success_rate': self.loading_stats.get('success_rate', 0)
            },
            'method_usage': {
                'advanced_parsing': self.loading_stats['advanced_parsing_used'],
                'fallback_processing': self.loading_stats['fallback_used']
            },
            'features_enabled': {
                'advanced_parsing': self.hybrid_processor.advanced_parsing_enabled,
                'image_extraction': self.hybrid_processor.extract_images,
                'structure_preservation': self.hybrid_processor.preserve_structure,
                'table_extraction': self.hybrid_processor.extract_tables,
                'hybrid_processing': self.hybrid_processor.hybrid_processing
            }
        }


# Utility functions for file operations
def scan_files_in_directory(directory, recursive=True):
    """
    Scan directory to get all files with detailed categorization
    
    Args:
        directory: Directory to scan
        recursive: Whether to scan recursively
    
    Returns:
        dict: Categorized file information
    """
    file_info = {
        'all_files': [],
        'docx_files': [],
        'doc_files': [],
        'pdf_files': [],
        'text_files': [],
        'image_files': [],
        'other_files': [],
        'total_size': 0,
        'file_count_by_ext': {}
    }
    
    try:
        if recursive:
            file_iterator = Path(directory).rglob('*')
        else:
            file_iterator = Path(directory).glob('*')
        
        for file_path in file_iterator:
            if file_path.is_file():
                file_str = str(file_path)
                file_ext = file_path.suffix.lower()
                file_size = file_path.stat().st_size
                
                file_info['all_files'].append(file_str)
                file_info['total_size'] += file_size
                
                # Count by extension
                if file_ext in file_info['file_count_by_ext']:
                    file_info['file_count_by_ext'][file_ext] += 1
                else:
                    file_info['file_count_by_ext'][file_ext] = 1
                
                # Categorize files
                if file_ext == '.docx':
                    file_info['docx_files'].append(file_str)
                elif file_ext == '.doc':
                    file_info['doc_files'].append(file_str)
                elif file_ext == '.pdf':
                    file_info['pdf_files'].append(file_str)
                elif file_ext in ['.txt', '.md', '.rst', '.log']:
                    file_info['text_files'].append(file_str)
                elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif']:
                    file_info['image_files'].append(file_str)
                else:
                    file_info['other_files'].append(file_str)
    
    except Exception as e:
        print(f"ERROR: Failed to scan directory {directory}: {e}")
    
    return file_info


def normalize_file_path(file_path):
    """
    Normalize file path for comparison
    
    Args:
        file_path: File path to normalize
    
    Returns:
        str: Normalized file path
    """
    return os.path.normpath(os.path.abspath(file_path))


def validate_file_path(file_path):
    """
    Validate if file path exists and is readable
    
    Args:
        file_path: Path to validate
    
    Returns:
        tuple: (is_valid, error_message)
    """
    try:
        if not os.path.exists(file_path):
            return False, "File does not exist"
        
        if not os.path.isfile(file_path):
            return False, "Path is not a file"
        
        if not os.access(file_path, os.R_OK):
            return False, "File is not readable"
        
        return True, None
        
    except Exception as e:
        return False, f"Error accessing file: {e}"


def get_file_info(file_path):
    """
    Get detailed information about a file
    
    Args:
        file_path: Path to the file
    
    Returns:
        dict: File information including size, extension, etc.
    """
    try:
        path_obj = Path(file_path)
        stat_info = os.stat(file_path)
        
        return {
            'name': path_obj.name,
            'stem': path_obj.stem,
            'suffix': path_obj.suffix.lower(),
            'size': stat_info.st_size,
            'size_mb': stat_info.st_size / (1024 * 1024),
            'modified': stat_info.st_mtime,
            'is_text_file': path_obj.suffix.lower() in ['.txt', '.md', '.rst', '.log'],
            'is_document': path_obj.suffix.lower() in ['.pdf', '.docx', '.doc', '.rtf'],
            'is_word_document': path_obj.suffix.lower() in ['.docx', '.doc'],
            'is_image': path_obj.suffix.lower() in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif'],
            'supports_advanced_parsing': path_obj.suffix.lower() in ['.docx', '.doc']
        }
    except Exception as e:
        return {'error': str(e)}


def scan_directory_files(directory, recursive=True):
    """
    Scan directory and return comprehensive file statistics
    
    Args:
        directory: Directory to scan
        recursive: Whether to scan recursively
    
    Returns:
        dict: Comprehensive directory statistics
    """
    try:
        stats = {
            'total_files': 0,
            'text_files': 0,
            'document_files': 0,
            'word_documents': 0,
            'docx_files': 0,
            'doc_files': 0,
            'image_files': 0,
            'pdf_files': 0,
            'other_files': 0,
            'total_size': 0,
            'large_files': [],
            'problematic_files': [],
            'advanced_parsing_candidates': 0,
            'file_extensions': {}
        }
        
        if recursive:
            file_iterator = Path(directory).rglob('*')
        else:
            file_iterator = Path(directory).glob('*')
        
        for file_path in file_iterator:
            if file_path.is_file():
                stats['total_files'] += 1
                
                file_info = get_file_info(file_path)
                if 'error' in file_info:
                    stats['problematic_files'].append(str(file_path))
                    continue
                
                stats['total_size'] += file_info['size']
                
                # Count by extension
                ext = file_info['suffix']
                if ext in stats['file_extensions']:
                    stats['file_extensions'][ext] += 1
                else:
                    stats['file_extensions'][ext] = 1
                
                # Categorize files
                if file_info['is_text_file']:
                    stats['text_files'] += 1
                elif file_info['is_document']:
                    stats['document_files'] += 1
                    if file_info['is_word_document']:
                        stats['word_documents'] += 1
                        if ext == '.docx':
                            stats['docx_files'] += 1
                        elif ext == '.doc':
                            stats['doc_files'] += 1
                    elif ext == '.pdf':
                        stats['pdf_files'] += 1
                elif file_info['is_image']:
                    stats['image_files'] += 1
                else:
                    stats['other_files'] += 1
                
                # Count files that support advanced parsing
                if file_info['supports_advanced_parsing']:
                    stats['advanced_parsing_candidates'] += 1
                
                # Track large files (>10MB)
                if file_info['size_mb'] > 10:
                    stats['large_files'].append({
                        'path': str(file_path),
                        'size_mb': file_info['size_mb']
                    })
        
        stats['total_size_mb'] = stats['total_size'] / (1024 * 1024)
        return stats
        
    except Exception as e:
        return {'error': str(e)}


def create_safe_reader(documents_dir, recursive=True, config=None):
    """
    Create an AdvancedDirectoryLoader instance with enhanced capabilities
    
    Args:
        documents_dir: Directory to read from
        recursive: Whether to read recursively
        config: Configuration object for advanced features
    
    Returns:
        AdvancedDirectoryLoader: Advanced loader instance
    """
    return AdvancedDirectoryLoader(
        input_dir=documents_dir,
        recursive=recursive,
        config=config
    )


def print_advanced_parsing_info():
    """Print information about advanced parsing capabilities"""
    print("\n=== ADVANCED DOCUMENT PARSING CAPABILITIES ===")
    
    # Check availability of advanced features
    features = [
        ("DOCX Advanced Parsing", DOCX_AVAILABLE, "python-docx"),
        ("Legacy DOC Conversion", PANDOC_AVAILABLE, "pypandoc"), 
        ("Image Processing", PIL_AVAILABLE, "pillow")
    ]
    
    for feature_name, available, package in features:
        status = "? AVAILABLE" if available else "? MISSING"
        print(f"  {feature_name:<25}: {status}")
        if not available:
            print(f"    Install with: pip install {package}")
    
    print(f"\nAdvanced Features:")
    print(f"  ?? DOCX Structure Preservation: Extract headings, tables, lists")
    print(f"  ???  Image Extraction from Documents: Extract and OCR embedded images") 
    print(f"  ?? Table Content Extraction: Parse tables with proper formatting")
    print(f"  ?? Legacy DOC Support: Convert .doc files using pandoc")
    print(f"  ?? Hybrid Processing: Combine text extraction with image OCR")
    print(f"  ?? Enhanced Text Cleaning: Remove null bytes and problematic characters")
    print("=" * 60)


# Test function for advanced parsing
def test_advanced_document_parsing(test_directory="./data/test_docs"):
    """
    Test function for advanced document parsing capabilities
    
    Args:
        test_directory: Directory containing test documents
    """
    print("?? Testing Advanced Document Parsing")
    
    if not os.path.exists(test_directory):
        print(f"WARNING: Test directory {test_directory} does not exist")
        return
    
    # Create test loader
    loader = create_safe_reader(test_directory, recursive=True)
    
    # Load documents
    documents, stats, failed_files = loader.load_data()
    
    # Print results
    print(f"\n?? Test Results:")
    summary = loader.get_processing_summary()
    
    print(f"  ?? Files found: {summary['total_files_found']}")
    print(f"    - DOCX files: {summary['file_breakdown']['docx_files']}")
    print(f"    - DOC files: {summary['file_breakdown']['doc_files']}")  
    print(f"    - Other files: {summary['file_breakdown']['other_files']}")
    
    print(f"  ? Documents created: {summary['processing_results']['documents_created']}")
    print(f"  ???  Images extracted: {summary['processing_results']['images_extracted']}")
    print(f"  ? Processing errors: {summary['processing_results']['processing_errors']}")
    print(f"  ?? Success rate: {summary['processing_results']['success_rate']:.1f}%")
    
    print(f"  ??  Advanced parsing used: {summary['method_usage']['advanced_parsing']}")
    print(f"  ?? Fallback processing used: {summary['method_usage']['fallback_processing']}")
    
    # Show sample document info
    if documents:
        print(f"\n?? Sample Document Info:")
        sample_doc = documents[0]
        print(f"  File: {sample_doc.metadata.get('file_name', 'Unknown')}")
        print(f"  Type: {sample_doc.metadata.get('file_type', 'Unknown')}")
        print(f"  Content length: {len(sample_doc.text)} characters")
        print(f"  Extraction method: {sample_doc.metadata.get('extraction_method', 'Unknown')}")
        
        # Show parsing info if available
        if 'parsing_info' in sample_doc.metadata:
            parsing_info = sample_doc.metadata['parsing_info']
            if 'structure_metadata' in parsing_info:
                structure = parsing_info['structure_metadata']
                print(f"  Structure info:")
                print(f"    - Headings: {structure.get('heading_count', 0)}")
                print(f"    - Paragraphs: {structure.get('paragraph_count', 0)}")
                print(f"    - Tables found: {parsing_info.get('tables_found', 0)}")
    
    print("?? Advanced parsing test completed\n")