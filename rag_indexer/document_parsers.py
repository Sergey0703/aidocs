#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Document parsers module for RAG Document Indexer
Specialized parsers for different document types
Enhanced with modular PDF processing architecture
"""

import os
import io
import zipfile
import tempfile
from pathlib import Path
from datetime import datetime
from llama_index.core import Document

# Advanced document parsing imports
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("INFO: python-docx not available. Install with: pip install python-docx")

try:
    import pypandoc
    PANDOC_AVAILABLE = True
except ImportError:
    PANDOC_AVAILABLE = False
    print("INFO: pypandoc not available. Install with: pip install pypandoc")

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

# Import enhanced PDF processor from separate module
try:
    from enhanced_pdf_processor import create_enhanced_pdf_processor
    ENHANCED_PDF_AVAILABLE = True
except ImportError:
    ENHANCED_PDF_AVAILABLE = False
    print("WARNING: Enhanced PDF processor not available. PDF processing will be limited.")

# Import utility functions
from file_utils_core import clean_content_from_null_bytes, clean_metadata_recursive, safe_read_file


class AdvancedDocxParser:
    """Advanced parser for .docx files with image extraction and structure preservation"""
    
    def __init__(self, extract_images=True, preserve_structure=True, extract_tables=True):
        """
        Initialize advanced DOCX parser
        
        Args:
            extract_images: Whether to extract images from documents
            preserve_structure: Whether to preserve document structure (headers, etc.)
            extract_tables: Whether to extract table content
        """
        self.extract_images = extract_images
        self.preserve_structure = preserve_structure
        self.extract_tables = extract_tables
        self.is_available = DOCX_AVAILABLE and PIL_AVAILABLE
        
        if not self.is_available:
            print("WARNING: Advanced DOCX parsing not available. Missing dependencies.")
    
    def extract_images_from_docx(self, docx_file_path):
        """
        Extract all images from DOCX file
        
        Args:
            docx_file_path: Path to DOCX file
        
        Returns:
            list: List of (image_data, image_name, image_format) tuples
        """
        if not self.is_available or not self.extract_images:
            return []
        
        images = []
        
        try:
            # Open DOCX as ZIP file to access images
            with zipfile.ZipFile(docx_file_path, 'r') as docx_zip:
                # Find all image files in the ZIP
                image_files = [f for f in docx_zip.namelist() 
                              if f.startswith('word/media/') and 
                              any(f.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp'])]
                
                for i, image_file in enumerate(image_files):
                    try:
                        # Extract image data
                        image_data = docx_zip.read(image_file)
                        
                        # Determine image format
                        image_name = os.path.basename(image_file)
                        image_format = os.path.splitext(image_name)[1].lower()
                        
                        # Create a more descriptive name
                        base_name = os.path.splitext(os.path.basename(docx_file_path))[0]
                        enhanced_name = f"{base_name}_image_{i+1}{image_format}"
                        
                        images.append((image_data, enhanced_name, image_format))
                        
                    except Exception as e:
                        print(f"   WARNING: Failed to extract image {image_file}: {e}")
                
            if images:
                print(f"   INFO: Extracted {len(images)} images from {os.path.basename(docx_file_path)}")
                
        except Exception as e:
            print(f"   ERROR: Failed to extract images from {docx_file_path}: {e}")
        
        return images
    
    def extract_table_content(self, doc):
        """
        Extract content from all tables in the document
        
        Args:
            doc: python-docx Document object
        
        Returns:
            str: Formatted table content
        """
        if not self.extract_tables:
            return ""
        
        table_content = []
        
        try:
            for i, table in enumerate(doc.tables):
                table_text = f"\n--- Table {i+1} ---\n"
                
                for row in table.rows:
                    row_cells = []
                    for cell in row.cells:
                        # Clean cell text and remove extra whitespace
                        cell_text = ' '.join(cell.text.strip().split())
                        row_cells.append(cell_text)
                    
                    # Join cells with tab separator
                    table_text += '\t'.join(row_cells) + '\n'
                
                table_text += "--- End Table ---\n"
                table_content.append(table_text)
                
        except Exception as e:
            print(f"   WARNING: Failed to extract table content: {e}")
        
        return '\n'.join(table_content)
    
    def extract_structured_content(self, doc):
        """
        Extract content while preserving document structure
        
        Args:
            doc: python-docx Document object
        
        Returns:
            tuple: (structured_text, metadata_info)
        """
        if not self.preserve_structure:
            # Simple text extraction
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            return '\n'.join(full_text), {}
        
        structured_content = []
        metadata_info = {
            'headings': [],
            'paragraph_count': 0,
            'heading_count': 0,
            'list_count': 0
        }
        
        try:
            for paragraph in doc.paragraphs:
                if not paragraph.text.strip():
                    continue
                
                # Check if paragraph is a heading
                if paragraph.style.name.startswith('Heading'):
                    level = paragraph.style.name.replace('Heading ', '')
                    heading_marker = '#' * min(int(level) if level.isdigit() else 1, 6)
                    structured_content.append(f"{heading_marker} {paragraph.text.strip()}")
                    
                    metadata_info['headings'].append({
                        'level': level,
                        'text': paragraph.text.strip()
                    })
                    metadata_info['heading_count'] += 1
                
                # Check if paragraph is a list item
                elif paragraph.style.name.startswith('List'):
                    structured_content.append(f"‚Ä¢ {paragraph.text.strip()}")
                    metadata_info['list_count'] += 1
                
                # Regular paragraph
                else:
                    structured_content.append(paragraph.text.strip())
                    metadata_info['paragraph_count'] += 1
            
        except Exception as e:
            print(f"   WARNING: Failed to extract structured content: {e}")
            # Fallback to simple extraction
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            return '\n'.join(full_text), {'extraction_fallback': True}
        
        return '\n'.join(structured_content), metadata_info
    
    def parse_docx_file(self, file_path):
        """
        Parse DOCX file with advanced features
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            tuple: (main_document, extracted_images, parsing_info)
        """
        if not self.is_available:
            print(f"   WARNING: Advanced DOCX parsing not available for {os.path.basename(file_path)}")
            return None, [], {'error': 'dependencies_missing'}
        
        try:
            print(f"   INFO: Advanced parsing of {os.path.basename(file_path)}")
            
            # Open document
            doc = DocxDocument(file_path)
            
            # Extract structured text content
            structured_text, structure_metadata = self.extract_structured_content(doc)
            
            # Extract table content
            table_content = self.extract_table_content(doc)
            
            # Combine text and table content
            full_text = structured_text
            if table_content:
                full_text += f"\n\n{table_content}"
            
            # Clean text
            full_text = clean_content_from_null_bytes(full_text)
            
            # Extract images
            extracted_images = self.extract_images_from_docx(file_path)
            
            # Create enhanced metadata
            parsing_info = {
                'parser_type': 'advanced_docx',
                'structure_preserved': self.preserve_structure,
                'images_extracted': len(extracted_images),
                'tables_found': len(doc.tables) if hasattr(doc, 'tables') else 0,
                'structure_metadata': structure_metadata,
                'extraction_features': {
                    'extract_images': self.extract_images,
                    'preserve_structure': self.preserve_structure,
                    'extract_tables': self.extract_tables
                }
            }
            
            # Create main document
            clean_file_path = clean_content_from_null_bytes(str(file_path))
            clean_file_name = clean_content_from_null_bytes(os.path.basename(file_path))
            
            main_metadata = {
                'file_path': clean_file_path,
                'file_name': clean_file_name,
                'file_type': 'docx',
                'file_size': os.path.getsize(file_path),
                'parsing_info': parsing_info,
                'content_length': len(full_text),
                'extraction_method': 'advanced_docx_parser'
            }
            
            # Clean metadata
            clean_main_metadata = clean_metadata_recursive(main_metadata)
            
            main_document = Document(
                text=full_text,
                metadata=clean_main_metadata
            )
            
            print(f"   SUCCESS: Extracted {len(full_text)} characters, {len(extracted_images)} images")
            
            return main_document, extracted_images, parsing_info
            
        except Exception as e:
            print(f"   ERROR: Failed to parse DOCX file {file_path}: {e}")
            return None, [], {'error': str(e), 'parser_type': 'advanced_docx'}


class LegacyDocConverter:
    """Converter for legacy .doc files using pandoc or other methods"""
    
    def __init__(self):
        """Initialize legacy DOC converter"""
        self.pandoc_available = PANDOC_AVAILABLE
        
        if not self.pandoc_available:
            print("INFO: pandoc not available for .doc conversion. .doc files will use fallback methods.")
    
    def convert_doc_to_text(self, doc_file_path):
        """
        Convert .doc file to text using available methods
        
        Args:
            doc_file_path: Path to .doc file
        
        Returns:
            tuple: (text_content, conversion_info)
        """
        conversion_info = {'method': 'none', 'success': False}
        
        try:
            # Method 1: Try pandoc if available
            if self.pandoc_available:
                try:
                    text_content = pypandoc.convert_file(doc_file_path, 'plain')
                    text_content = clean_content_from_null_bytes(text_content)
                    
                    if text_content and text_content.strip():
                        conversion_info = {
                            'method': 'pandoc',
                            'success': True,
                            'content_length': len(text_content)
                        }
                        print(f"   INFO: Converted .doc using pandoc: {len(text_content)} characters")
                        return text_content, conversion_info
                
                except Exception as e:
                    print(f"   WARNING: Pandoc conversion failed: {e}")
            
            # Method 2: Try to read as binary and extract readable text (very basic)
            try:
                with open(doc_file_path, 'rb') as f:
                    raw_content = f.read()
                
                # Very basic text extraction from binary data
                # This is a fallback method and may not work well
                text_content = raw_content.decode('latin-1', errors='ignore')
                
                # Filter out non-printable characters and keep only readable text
                readable_chars = []
                for char in text_content:
                    if char.isprintable() or char.isspace():
                        readable_chars.append(char)
                
                filtered_content = ''.join(readable_chars)
                
                # Remove excessive whitespace and clean up
                lines = [line.strip() for line in filtered_content.split('\n') if line.strip()]
                final_content = '\n'.join(lines)
                
                # Clean null bytes
                final_content = clean_content_from_null_bytes(final_content)
                
                if final_content and len(final_content) > 100:  # Minimum reasonable length
                    conversion_info = {
                        'method': 'binary_extraction',
                        'success': True,
                        'content_length': len(final_content),
                        'warning': 'Low quality extraction method'
                    }
                    print(f"   WARNING: Using basic binary extraction for .doc file: {len(final_content)} characters")
                    return final_content, conversion_info
                
            except Exception as e:
                print(f"   WARNING: Binary extraction failed: {e}")
            
            # If all methods fail
            conversion_info = {
                'method': 'failed',
                'success': False,
                'error': 'No conversion method succeeded'
            }
            
            return "", conversion_info
            
        except Exception as e:
            conversion_info = {
                'method': 'error',
                'success': False,
                'error': str(e)
            }
            return "", conversion_info


class HybridDocumentProcessor:
    """
    Processor that combines text extraction with image OCR for complete document processing
    Now uses modular PDF processor architecture
    """
    
    def __init__(self, config=None):
        """
        Initialize hybrid document processor
        
        Args:
            config: Configuration object with processing settings
        """
        self.config = config
        
        # Load settings from config
        if config:
            doc_settings = config.get_document_parsing_settings()
            self.advanced_parsing_enabled = doc_settings.get('advanced_parsing_enabled', True)
            self.extract_images = doc_settings.get('extract_images', True)
            self.preserve_structure = doc_settings.get('preserve_structure', True)
            self.extract_tables = doc_settings.get('extract_tables', True)
            self.hybrid_processing = doc_settings.get('hybrid_processing', True)
            self.combine_results = doc_settings.get('combine_results', True)
            self.image_quality = doc_settings.get('image_quality', 'high')
        else:
            # Default settings
            self.advanced_parsing_enabled = True
            self.extract_images = True
            self.preserve_structure = True
            self.extract_tables = True
            self.hybrid_processing = True
            self.combine_results = True
            self.image_quality = 'high'
        
        # Initialize specialized parsers
        self.docx_parser = AdvancedDocxParser(
            extract_images=self.extract_images,
            preserve_structure=self.preserve_structure,
            extract_tables=self.extract_tables
        ) if self.advanced_parsing_enabled else None
        
        self.doc_converter = LegacyDocConverter() if self.advanced_parsing_enabled else None
        
        # Initialize enhanced PDF processor from separate module
        if self.advanced_parsing_enabled and ENHANCED_PDF_AVAILABLE:
            try:
                self.pdf_processor = create_enhanced_pdf_processor(config)
                print(f"   Ì†ΩÌ≥Ñ Enhanced PDF processor initialized")
            except Exception as e:
                print(f"   WARNING: Enhanced PDF processor initialization failed: {e}")
                self.pdf_processor = None
        else:
            self.pdf_processor = None
            if not ENHANCED_PDF_AVAILABLE:
                print(f"   WARNING: Enhanced PDF processor not available - install enhanced_pdf_processor.py")
        
        # OCR processor will be injected when needed
        self.ocr_processor = None
    
    def set_ocr_processor(self, ocr_processor):
        """
        Set OCR processor for image processing
        
        Args:
            ocr_processor: OCR processor instance
        """
        self.ocr_processor = ocr_processor
        
        # Also set for PDF processor if available
        if self.pdf_processor:
            self.pdf_processor.set_ocr_processor(ocr_processor)
            print(f"   Ì†æÌ¥ñ OCR processor integrated with PDF processor")
    
    def process_pdf_file(self, file_path):
        """
        Process PDF file with enhanced PDF processor
        
        Args:
            file_path: Path to PDF file
        
        Returns:
            list: List of Document objects
        """
        if not self.pdf_processor or not self.advanced_parsing_enabled:
            # Fallback to simple processing
            print(f"   WARNING: Enhanced PDF processing not available for {os.path.basename(file_path)}")
            return self._simple_file_processing(file_path)
        
        try:
            # Use enhanced PDF processor from separate module
            documents = self.pdf_processor.process_pdf_file(file_path)
            return documents if documents else []
            
        except Exception as e:
            print(f"   ERROR: Enhanced PDF processing failed for {file_path}: {e}")
            print(f"   INFO: Falling back to simple processing...")
            return self._simple_file_processing(file_path)
    
    def process_docx_file(self, file_path):
        """
        Process DOCX file with advanced parsing and image extraction
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            list: List of Document objects (main document + image documents)
        """
        if not self.docx_parser or not self.advanced_parsing_enabled:
            # Fallback to simple processing
            return self._simple_file_processing(file_path)
        
        try:
            # Parse DOCX with advanced features
            main_document, extracted_images, parsing_info = self.docx_parser.parse_docx_file(file_path)
            
            documents = []
            
            # Add main document if successful
            if main_document:
                documents.append(main_document)
            
            # Process extracted images with OCR if enabled
            if extracted_images and self.hybrid_processing and self.ocr_processor:
                print(f"   INFO: Processing {len(extracted_images)} extracted images with OCR")
                
                for image_data, image_name, image_format in extracted_images:
                    try:
                        # Create temporary file for OCR processing
                        with tempfile.NamedTemporaryFile(suffix=image_format, delete=False) as temp_file:
                            temp_file.write(image_data)
                            temp_file_path = temp_file.name
                        
                        # Process with OCR
                        ocr_document = self.ocr_processor.process_single_image(temp_file_path)
                        
                        if ocr_document:
                            # Enhance metadata to indicate it's from a document
                            ocr_document.metadata.update({
                                'source_document': os.path.basename(file_path),
                                'extraction_method': 'docx_image_ocr',
                                'original_image_name': image_name
                            })
                            documents.append(ocr_document)
                        
                        # Clean up temporary file
                        os.unlink(temp_file_path)
                        
                    except Exception as e:
                        print(f"   WARNING: Failed to process extracted image {image_name}: {e}")
            
            return documents
            
        except Exception as e:
            print(f"   ERROR: Advanced DOCX processing failed for {file_path}: {e}")
            return self._simple_file_processing(file_path)
    
    def process_doc_file(self, file_path):
        """
        Process legacy .doc file
        
        Args:
            file_path: Path to .doc file
        
        Returns:
            list: List of Document objects
        """
        if not self.doc_converter or not self.advanced_parsing_enabled:
            return self._simple_file_processing(file_path)
        
        try:
            # Convert .doc to text
            text_content, conversion_info = self.doc_converter.convert_doc_to_text(file_path)
            
            if conversion_info['success'] and text_content:
                # Create document
                clean_file_path = clean_content_from_null_bytes(str(file_path))
                clean_file_name = clean_content_from_null_bytes(os.path.basename(file_path))
                
                metadata = {
                    'file_path': clean_file_path,
                    'file_name': clean_file_name,
                    'file_type': 'doc',
                    'file_size': os.path.getsize(file_path),
                    'conversion_info': conversion_info,
                    'content_length': len(text_content),
                    'extraction_method': 'legacy_doc_converter'
                }
                
                # Clean metadata
                clean_metadata = clean_metadata_recursive(metadata)
                
                document = Document(
                    text=text_content,
                    metadata=clean_metadata
                )
                
                return [document]
            else:
                print(f"   WARNING: Failed to convert .doc file: {file_path}")
                return self._simple_file_processing(file_path)
                
        except Exception as e:
            print(f"   ERROR: Legacy .doc processing failed for {file_path}: {e}")
            return self._simple_file_processing(file_path)
    
    def _simple_file_processing(self, file_path):
        """
        Fallback to simple file processing
        
        Args:
            file_path: Path to file
        
        Returns:
            list: List of Document objects
        """
        try:
            content, error_code = safe_read_file(file_path)
            
            if content:
                clean_file_path = clean_content_from_null_bytes(str(file_path))
                clean_file_name = clean_content_from_null_bytes(os.path.basename(file_path))
                
                metadata = {
                    'file_path': clean_file_path,
                    'file_name': clean_file_name,
                    'file_type': os.path.splitext(file_path)[1].lower(),
                    'file_size': os.path.getsize(file_path),
                    'extraction_method': 'simple_text_reader',
                    'content_length': len(content)
                }
                
                if error_code:
                    metadata['encoding_fallback'] = error_code
                
                # Clean metadata
                clean_metadata = clean_metadata_recursive(metadata)
                
                document = Document(
                    text=content,
                    metadata=clean_metadata
                )
                
                return [document]
            else:
                return []
                
        except Exception as e:
            print(f"   ERROR: Simple file processing failed for {file_path}: {e}")
            return []
    
    def get_processing_capabilities(self):
        """
        Get information about available processing capabilities
        
        Returns:
            dict: Capability information
        """
        return {
            'advanced_parsing_enabled': self.advanced_parsing_enabled,
            'docx_parser_available': self.docx_parser is not None,
            'doc_converter_available': self.doc_converter is not None,
            'pdf_processor_available': self.pdf_processor is not None,
            'ocr_processor_available': self.ocr_processor is not None,
            'enhanced_pdf_available': ENHANCED_PDF_AVAILABLE,
            'features': {
                'extract_images': self.extract_images,
                'preserve_structure': self.preserve_structure,
                'extract_tables': self.extract_tables,
                'hybrid_processing': self.hybrid_processing,
                'combine_results': self.combine_results
            }
        }
    
    def print_capabilities_summary(self):
        """Print summary of processing capabilities"""
        capabilities = self.get_processing_capabilities()
        
        print(f"\nÌ†ΩÌ≥Ñ Hybrid Document Processor Capabilities:")
        print(f"   Ì†ΩÌ≥ã Advanced parsing: {'‚úÖ' if capabilities['advanced_parsing_enabled'] else '‚ùå'}")
        print(f"   Ì†ΩÌ≥Ñ DOCX parser: {'‚úÖ' if capabilities['docx_parser_available'] else '‚ùå'}")
        print(f"   Ì†ΩÌ≥Ñ DOC converter: {'‚úÖ' if capabilities['doc_converter_available'] else '‚ùå'}")
        print(f"   Ì†ΩÌ≥Ñ Enhanced PDF processor: {'‚úÖ' if capabilities['pdf_processor_available'] else '‚ùå'}")
        print(f"   Ì†æÌ¥ñ OCR processor: {'‚úÖ' if capabilities['ocr_processor_available'] else '‚ùå'}")
        
        features = capabilities['features']
        print(f"\n‚öôÔ∏è Feature Settings:")
        for feature, enabled in features.items():
            status = "‚úÖ" if enabled else "‚ùå"
            feature_name = feature.replace('_', ' ').title()
            print(f"   {status} {feature_name}")
        
        if not capabilities['enhanced_pdf_available']:
            print(f"\n‚ö†Ô∏è Enhanced PDF processing not available:")
            print(f"   Install enhanced_pdf_processor.py module for full PDF support")


# Factory functions for creating processor instances
def create_hybrid_document_processor(config=None):
    """
    Create hybrid document processor instance
    
    Args:
        config: Configuration object
    
    Returns:
        HybridDocumentProcessor: Configured processor
    """
    return HybridDocumentProcessor(config)


def create_docx_parser(extract_images=True, preserve_structure=True, extract_tables=True):
    """
    Create advanced DOCX parser instance
    
    Args:
        extract_images: Whether to extract images
        preserve_structure: Whether to preserve structure
        extract_tables: Whether to extract tables
    
    Returns:
        AdvancedDocxParser: Configured parser
    """
    return AdvancedDocxParser(extract_images, preserve_structure, extract_tables)


def create_doc_converter():
    """
    Create legacy DOC converter instance
    
    Returns:
        LegacyDocConverter: DOC converter
    """
    return LegacyDocConverter()


def check_document_processing_capabilities():
    """
    Check available document processing capabilities
    
    Returns:
        dict: Comprehensive capability information
    """
    capabilities = {
        'docx_available': DOCX_AVAILABLE,
        'pandoc_available': PANDOC_AVAILABLE,
        'pil_available': PIL_AVAILABLE,
        'enhanced_pdf_available': ENHANCED_PDF_AVAILABLE,
        'overall_status': 'unknown'
    }
    
    # Determine overall status
    basic_available = capabilities['docx_available'] or capabilities['pandoc_available']
    
    if basic_available and capabilities['enhanced_pdf_available']:
        capabilities['overall_status'] = 'excellent'
    elif basic_available:
        capabilities['overall_status'] = 'good'
    elif capabilities['enhanced_pdf_available']:
        capabilities['overall_status'] = 'pdf_only'
    else:
        capabilities['overall_status'] = 'limited'
    
    # Feature availability
    capabilities['features'] = {
        'docx_parsing': capabilities['docx_available'] and capabilities['pil_available'],
        'doc_conversion': capabilities['pandoc_available'],
        'image_extraction': capabilities['pil_available'],
        'pdf_processing': capabilities['enhanced_pdf_available'],
        'structure_preservation': capabilities['docx_available'],
        'table_extraction': capabilities['docx_available']
    }
    
    return capabilities


def print_document_processing_status():
    """Print comprehensive document processing status"""
    capabilities = check_document_processing_capabilities()
    
    print("Ì†ΩÌ≥Ñ Document Processing Capabilities:")
    print(f"   DOCX parsing: {'‚úÖ' if capabilities['docx_available'] else '‚ùå'}")
    print(f"   DOC conversion: {'‚úÖ' if capabilities['pandoc_available'] else '‚ùå'}")
    print(f"   Image processing: {'‚úÖ' if capabilities['pil_available'] else '‚ùå'}")
    print(f"   Enhanced PDF: {'‚úÖ' if capabilities['enhanced_pdf_available'] else '‚ùå'}")
    
    print(f"\nÌ†ΩÌ∫Ä Available Features:")
    for feature, available in capabilities['features'].items():
        status = "‚úÖ" if available else "‚ùå"
        feature_name = feature.replace('_', ' ').title()
        print(f"   {status} {feature_name}")
    
    print(f"\nÌ†ΩÌ≥ä Overall Status: {capabilities['overall_status'].upper()}")
    
    # Recommendations
    recommendations = []
    if not capabilities['docx_available']:
        recommendations.append("Install python-docx: pip install python-docx")
    if not capabilities['pandoc_available']:
        recommendations.append("Install pypandoc: pip install pypandoc")
    if not capabilities['pil_available']:
        recommendations.append("Install Pillow: pip install Pillow")
    if not capabilities['enhanced_pdf_available']:
        recommendations.append("Install enhanced_pdf_processor.py module for PDF support")
    
    if recommendations:
        print(f"\nÌ†ΩÌ≤° Recommendations:")
        for rec in recommendations:
            print(f"   ‚Ä¢ {rec}")
    
    return capabilities


if __name__ == "__main__":
    # Test document processing capabilities when run directly
    print("Ì†æÌ∑™ Document Parsers - Capability Test")
    print("=" * 60)
    
    capabilities = print_document_processing_status()
    
    # Test hybrid processor creation
    print(f"\nÌ†ΩÌ¥ß Testing Hybrid Processor Creation...")
    try:
        processor = create_hybrid_document_processor()
        processor.print_capabilities_summary()
        print(f"‚úÖ Hybrid processor created successfully")
    except Exception as e:
        print(f"‚ùå Hybrid processor creation failed: {e}")
    
    print("=" * 60)